"""
Document Ingestion Module
Handles reading files, chunking text, creating embeddings, and storing in Pinecone
"""
import os
from typing import List
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain_community.embeddings import HuggingFaceEmbeddings
from pinecone import Pinecone, ServerlessSpec
import config


class DocumentIngestion:
    def __init__(self):
        """Initialize document ingestion with embeddings and vector store"""
        self.embeddings = HuggingFaceEmbeddings(
            model_name=config.EMBEDDING_MODEL,
            model_kwargs={'device': 'cpu'}
        )
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.CHUNK_SIZE,
            chunk_overlap=config.CHUNK_OVERLAP,
            length_function=len,
        )
        
        # Initialize Pinecone
        self.pc = None
        self.index = None
        if config.PINECONE_API_KEY:
            self._init_pinecone()
    
    def _init_pinecone(self):
        """Initialize Pinecone connection and index"""
        try:
            self.pc = Pinecone(api_key=config.PINECONE_API_KEY)
            
            # Check if index exists, create if not
            if config.PINECONE_INDEX_NAME not in self.pc.list_indexes().names():
                self.pc.create_index(
                    name=config.PINECONE_INDEX_NAME,
                    dimension=config.EMBEDDING_DIMENSION,
                    metric='cosine',
                    spec=ServerlessSpec(
                        cloud='aws',
                        region=config.PINECONE_ENVIRONMENT or 'us-east-1'
                    )
                )
            
            self.index = self.pc.Index(config.PINECONE_INDEX_NAME)
            print(f"✅ Connected to Pinecone index: {config.PINECONE_INDEX_NAME}")
        except Exception as e:
            print(f"⚠️ Pinecone initialization error: {e}")
    
    def read_text_file(self, file_path: str) -> str:
        """Read content from a text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            print(f"Error reading text file: {e}")
            return ""
    
    def read_pdf_file(self, file_path: str) -> str:
        """Read content from a PDF file"""
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            print(f"Error reading PDF file: {e}")
            return ""
    
    def read_docx_file(self, file_path: str) -> str:
        """Read content from a DOCX file"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            print(f"Error reading DOCX file: {e}")
            return ""
    
    def read_file(self, file_path: str) -> str:
        """Read file based on extension"""
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.txt':
            return self.read_text_file(file_path)
        elif ext == '.pdf':
            return self.read_pdf_file(file_path)
        elif ext == '.docx':
            return self.read_docx_file(file_path)
        else:
            print(f"Unsupported file type: {ext}")
            return ""
    
    def chunk_text(self, text: str, metadata: dict = None) -> List[Document]:
        """Split text into chunks"""
        if metadata is None:
            metadata = {}
        
        documents = [Document(page_content=text, metadata=metadata)]
        chunks = self.text_splitter.split_documents(documents)
        return chunks
    
    def create_embeddings(self, texts: List[str]) -> List[List[float]]:
        """Create embeddings for texts"""
        return self.embeddings.embed_documents(texts)
    
    def store_in_pinecone(self, chunks: List[Document], file_name: str):
        """Store document chunks in Pinecone"""
        if not self.index:
            print("⚠️ Pinecone index not initialized. Please check your API key.")
            return
        
        try:
            # Prepare data for Pinecone
            texts = [chunk.page_content for chunk in chunks]
            embeddings = self.create_embeddings(texts)
            
            # Create vectors with metadata
            vectors = []
            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                vector_id = f"{file_name}_{i}"
                metadata = {
                    "text": chunk.page_content,
                    "source": file_name,
                    **chunk.metadata
                }
                vectors.append({
                    "id": vector_id,
                    "values": embedding,
                    "metadata": metadata
                })
            
            # Upsert to Pinecone in batches
            batch_size = 100
            for i in range(0, len(vectors), batch_size):
                batch = vectors[i:i + batch_size]
                self.index.upsert(vectors=batch)
            
            print(f"✅ Stored {len(vectors)} chunks from {file_name} in Pinecone")
        except Exception as e:
            print(f"❌ Error storing in Pinecone: {e}")
    
    def ingest_file(self, file_path: str, file_name: str = None):
        """Complete ingestion pipeline for a file"""
        if file_name is None:
            file_name = os.path.basename(file_path)
        
        print(f"📄 Processing {file_name}...")
        
        # Read file
        text = self.read_file(file_path)
        if not text:
            print(f"❌ No content extracted from {file_name}")
            return
        
        # Chunk text
        chunks = self.chunk_text(text, metadata={"source": file_name})
        print(f"📝 Created {len(chunks)} chunks")
        
        # Store in Pinecone
        self.store_in_pinecone(chunks, file_name)
        
        print(f"✅ Successfully ingested {file_name}")
